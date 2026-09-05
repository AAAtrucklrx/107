"""107 杯提交材料生成：作品简介 + 设计文档（docx 与 HTML 同源产出）。

设计文档按学校通知要求组织：设计思路、技术架构、功能模块、突出技术难度；
首节附评审四维度（创新性/实用性/技术难度/完成度）的评分项→章节→证据映射表；
嵌入 Mermaid 图（diagrams/*.png）、Prompt 工程节选与实测数据表。

用法：py submission/build_docs.py
PDF 由 Edge headless 将 HTML 打印输出（--print-to-pdf）。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent
DIAG_DIR = OUT_DIR / "diagrams"
PRIMARY = RGBColor(0x03, 0x4E, 0xA1)
MUTED = RGBColor(0x56, 0x66, 0x72)

# 条目类型：
#   str                          → 正文段落（开头"标签：内容"格式自动加粗标签）
#   ("h3", text)                 → 三级小节标题
#   ("img", filename, caption)   → 嵌图 + 图注
#   ("table", headers, rows)     → 表格
#   ("pre", text)                → 等宽代码块（Prompt 节选）


# ============================ 作品简介 ============================

BRIEF_TITLE = "小蜗——科大校园智能助手（作品简介）"
BRIEF_SUBTITLE = "一〇七杯算力与智能体开发大赛 · 智能体赛道 · 本科生队伍"

BRIEF_SECTIONS: list[tuple[str, list]] = [
    ("一、作品背景", [
        "先看一个日常场景：准备下学期选课时，要开公共查询确认培养方案必修缺口，去评课社区翻候选课的"
        "评分与评论，再对着课表逐门核对时间冲突，还要提防群里转发的过时通知——一次选课决策横跨至少三个系统、"
        "人工来回比对。",
        "中国科学技术大学学生的日常信息分散在十余个彼此独立的系统中：综合教务系统（成绩、课表、选课）、"
        "公共查询平台（空教室、培养计划）、评课社区 icourse.club（4.4 万条课程评价）、青春科大第二课堂平台（活动）、"
        "以及大量办事网页与公众号通知。查一件小事往往需要在多个系统间来回切换、手动比对。",
        "校园里已有的学生工具（导航站、门户类 App 等）本质上是\"人找信息\"的被动导航：没有 AI 统一入口，"
        "没有跨系统信息整合，也没有个性化与推荐能力。与此同时，群聊截图与过时通知泛滥，"
        "\"教务信息真假难辨、无从核验\"是学生的真实痛点。",
        "小蜗（Xiaowo）是面向科大师生的校园智能助手：用一句自然语言完成查、算、比、约、推等真实任务，"
        "并把\"答案来自哪里、是否可靠\"作为一等公民直接呈现给用户。",
    ]),
    ("二、系统工作流总览", [
        ("img", "2-qa-workflow.png",
         "图 1  小蜗统一 QA 工作流：前置分流 → 高置信确定性直连路由 / think 自主决策循环（22 条规则约束）"
         "→ 并行工具执行 → 增量流式合成（含降级路径）"),
        "上图即小蜗的核心：多类智能体角色（意图分类、决策、工具执行、证据核验、合成）在一张 LangGraph 图内协作；"
        "高置信问题跳过决策大模型直连工具，正文逐 token 流式推送（实测首字 3.1 秒）；"
        "大模型故障时按规则降级，永不假死。详见设计文档。",
    ]),
    ("三、解决的问题", [
        "多系统切换：一句话替代\"打开教务网—登录—找菜单—截图—再开公共查询比对\"的完整流程，"
        "覆盖成绩、GPA、课表、考试安排、空教室、日程、活动、办事流程等高频场景。",
        "信息可信度：回答遵循证据门槛——确定结论必须有审核通过的官方一手来源，或两个独立且一致的可靠来源；"
        "证据不足时明确说明，绝不编造。所有降级/缓存数据都带来源标识（实时数据/本地缓存/第三方工具）。",
        "选课信息过载：基于真实评课数据与个人培养方案的个性化推荐，课程范围是硬条件，"
        "兴趣/工作量等偏好软排序，并提供节次/周次级课表冲突检查。",
        "AI 幻觉与越权：个人数据只取当前登录用户经 CAS/成绩档案验证的身份，取不到就不猜；"
        "敏感请求固定拒绝；联网抓取经 SSRF 域白名单与速率限制。",
    ]),
    ("四、核心功能", [
        "智能问答（RAG）：80 篇校园知识文档、769 条向量分块的混合检索（向量 + BM25），回答附官方来源网址。",
        "联网证据通道：本地资料不足或问题要求最新信息时，经隐私清洗后联网检索（博查 Web Search API 国内直连），"
        "并按证据门槛核验后才输出确定结论；科大相关问题优先检索微信公众号官方号。",
        "课业助手：成绩、GPA、周课表（1–13 小节精确映射科大校历时间）、考试安排、空教室查询。",
        "今日卡弹窗：登录用户打开页面即自动弹出今日概览——今日课程（真实课表周次+星期匹配）与今日开始的活动"
        "（北京时区判定），无课/无活动分别如实标注；顺带解决了“打开页面不知道今天有什么”的高频入口问题。",
        "选课顾问：评课数据驱动的课程推荐（必修/方案内选修/方向补充）、教师对比、退补选压力评估、独立课表冲突检查。",
        "培养方案：登录后只读当前用户的专业与个人方案，展示修读进度与缺口。",
        "日程与活动：自然语言时间解析（\"明天下午3点到4点\"）、冲突检测、课表导入；"
        "青春科大实时活动查询与每日个性化推荐（紧迫度/课表空闲/兴趣/热度四因子）。",
        "官方入口与校园工具：强操作诉求给出已核实的官方入口跳转；社区工具经申请—审核—上架流程接入。",
        "管理后台：知识审核（逐块批准/拒绝）、校园工具治理、审计与 generation 发布/回滚。",
    ]),
    ("五、使用的模型及 API 调用方式", [
        "生成模型：科大 LLM 平台 deepseek-v4-flash（api.llm.ustc.edu.cn）。OpenAI 兼容协议，"
        "经 LangChain ChatOpenAI 封装，用于意图分类、think 决策、compose 回答合成、查询改写、证据抽取五类调用点。",
        "向量模型：科大 LLM 平台 qwen3-embedding，用于知识分块向量化与语义检索。",
        "图片 OCR：科大 LLM 平台 unlimited-ocr，用于微信公众号文章配图的文字提取。",
        "联网搜索：博查 Web Search API（国内直连，无需自托管搜索引擎，搜索源可配置切换）；"
        "页面抓取经 Crawl4AI sidecar（egress/robots/连接固定健康检查前置）。",
        "工具调用：30+ 内置工具经统一 Tool Registry 注册，智能体在决策循环中按需选择并可并行调用；"
        "个人数据类工具（成绩/课表/培养方案等）的学号由认证上下文强制注入，不接受用户输入伪造。",
        "校内服务对接：科大统一身份认证（CAS）、综合教务 jw API、青春科大 young 平台。",
    ]),
    ("六、创新点", [
        "证据门槛的可信回答：将\"AI 不编造\"做成机制而非提示词——确定结论须过一手权威来源或双独立来源门槛，"
        "未过门槛的回答降级为\"证据不足\"并如实告知，附可点击的官方来源。",
        "统一 QA 智能体架构：意图分类 → think 自主决策（≤4 轮，规则约束 + 意图裁剪工具表）→ 并行工具调用 → "
        "compose 合成的 LangGraph 流程；高置信问题（成绩/课表/考试等）走确定性直连路由跳过决策 LLM，"
        "配合 compose 增量流式，首字延迟实测约 3 秒。",
        "AI 输出可治理：公开网页证据进入\"采集→清洗→逐块审核→generation 发布→可回滚\"流水线，"
        "原始快照不可变，发布索引与回滚约束保证知识库变更全程可追溯。",
        "严格的身份与隐私边界：专业/年级只取已验证档案、匿名选择不继承、demo 与生产数据永久隔离；"
        "联网抓取不携带任何校内凭证，SSRF 防护覆盖内网/回环/云元数据/重定向。",
        "全链路流式与三层降级：SSE 逐 token 正文流式；LLM 平台故障时意图降级、工具摘要直出、熔断不传染，"
        "生成超时但已有流式正文时保留部分内容收尾，不整体失败。",
    ]),
    ("七、运行与部署", [
        "公网演示环境（比赛环境）：http://114.214.241.119:8850（competition + demo 模式，合成演示身份）。",
        "本地部署：pip install -r requirements.txt → frontend 内 npm ci && npm run build → "
        "复制 .env.example 为 .env → python init_check.py → uvicorn 启动，详见仓库 README。",
        "测试基线：pytest tests/web 225 项全绿；工具/节点/安全等 verify 系列脚本 260+ 断言通过。",
    ]),
]


# ============================ 设计文档 ============================

DESIGN_TITLE = "小蜗——科大校园智能助手（设计文档）"
DESIGN_SUBTITLE = "一〇七杯算力与智能体开发大赛 · 智能体赛道 · 本科生队伍"

DESIGN_SECTIONS: list[tuple[str, list]] = [
    ("一、评分项映射", [
        "本文档章节按学校通知要求组织：设计思路（二）、技术架构（三）、功能模块（四）、突出技术难度（五）。"
        "评审参考的评分项对应章节与核心证据如下，便于快速定位：",
        ("table",
         ["评分维度", "分项标准", "对应章节", "核心证据"],
         [
             ["创新性", "场景创新", "§2.1", "信息可信度与跨系统整合——未被现有工具覆盖的需求"],
             ["创新性", "交互与体验创新", "§2.3", "全链路流式（首字 3.1s）、结构化数据卡、思考过程可视、降级透明"],
             ["创新性", "多智能体协作创新", "§3.2 图 2", "五角色协作的 LangGraph 统一图 + 规则/LLM 混合编排 + 并行工具"],
             ["实用性", "场景真实性与问题匹配度", "§4.1", "四类真实校内数据源；教学/生活/服务三场景全覆盖"],
             ["实用性", "功能完整性与解决方案有效性", "§4.2", "十类功能模块全部实现并在生产环境运行"],
             ["实用性", "可扩展性与推广价值", "§4.3", "生态工具协议、配置驱动服务、架构可迁移其他高校"],
             ["技术难度", "Workflow 与系统设计", "§3 图 1", "五层架构、模块划分清晰、generation 版本化"],
             ["技术难度", "大模型应用深度", "§5.1 图 3", "混合检索/查询改写/语义缓存 + 22 条决策规则 Prompt + 证据门槛"],
             ["技术难度", "调试与优化能力", "§5.2 图 4", "实测优化数据表、三层降级、225 项回归测试"],
             ["完成度", "作品完整性", "§6", "生产部署 + 公网可访问 + 测试基线全绿"],
             ["完成度", "设计文档与材料质量", "§6", "仓库文档体系 + 本文档架构图与模型调用说明"],
             ["完成度", "演示效果", "§6", "公网演示环境 + ≤5 分钟演示视频"],
         ]),
    ]),
    ("二、设计思路", [
        ("h3", "2.1 问题定义：两个真实且未被解决的需求"),
        "先看一个真实场景：一名科大本科生准备下学期选课——他要打开公共查询看培养方案确认必修缺口，"
        "再开评课社区翻几门候选课的评分与评论，最后对着课表逐门核对时间冲突；期间还要留意群里"
        "转发的选课通知截图，真假难辨。一次选课决策要横跨至少三个系统、人工来回比对。",
        "备赛阶段完整调研了科大 8 个官方平台、12 个学生工具与历届获奖项目，结论：现有工具全部是"
        "\"人找信息\"的被动模式——导航站聚合链接、门户 App 展示信息，没有一个同时解决两个更根本的问题："
        "(1) 跨系统任务（如上例的选课决策）只能人工比对；(2) 信息可信度无从核验——群聊截图与过时通知泛滥，"
        "学生无法确认一条教务信息是真是假、是不是最新。",
        "小蜗的设计正是围绕这两个\"没人做\"的需求：跨系统任务一句话完成 + 每个答案可核验来源。",
        ("h3", "2.2 设计目标与核心取舍"),
        "设计目标：一句自然语言完成真实任务（查、算、比、约、推），且每个答案可核验来源、可区分实时与缓存、"
        "可明确\"证据不足\"。",
        "核心取舍一：可靠性优先于功能数量——宁可回答\"暂无数据\"也不编造；\"可信度\"作为一等公民贯穿全系统"
        "（证据门槛、来源标识、降级标注），而非事后补丁。",
        "核心取舍二：规则与 LLM 自主决策混合编排——高置信意图走确定性路由完全跳过决策 LLM（快且行为可预期），"
        "长尾意图交给带 22 条规则约束的 LLM 自主决策（覆盖面广）。纯规则系统无法覆盖长尾，纯 LLM 决策又慢"
        "且不可控，两层结合兼得可靠性与灵活性（详见 §3.2）。",
        "核心取舍三：身份与数据边界从严——个人数据只取当前登录用户经 CAS/成绩档案验证的身份，取不到就不猜、"
        "不继承匿名选择；demo 与生产数据永久隔离。",
        ("h3", "2.3 差异化定位与交互设计"),
        "差异化定位：不做通用聊天助手，做\"科大专用工具 + 个人学业数据 + 严格证据门槛\"的结合体，"
        "这是与通用大模型对话产品和学生自建导航站的根本区别。",
        "全链路流式：从用户发送问题到看到第一个字，实测 3.1 秒；正文逐 token 流式推送（answer.delta 事件，"
        "约 16 字一批），工具产出的结构化数据（成绩单/课表/推荐列表）以数据卡先于正文呈现，用户先看数据再听解读。",
        "今日卡弹窗（2026-09-05）：登录用户打开页面即呈现今日课程与今日活动，把“打开页面不知道今天有什么”的"
        "高频问题在零操作内解决；与全站弹窗统一 Radix 交互（Esc/遮罩/焦点管理），加载/失败/无课三态如实区分，"
        "活动窗口严格锚定北京时区，不把“可报名”误说成“今天开始”。",
        "过程透明：智能体的思考决策以折叠卡形式实时展示（thought.step 事件）——用户能看到\"小蜗决定查成绩工具\""
        "这类动作播报，而不是黑盒等待；所有降级/缓存数据带来源标识（实时/本地缓存/第三方）。",
        "多形态身份适配：匿名可问公共知识、Demo 演示身份展示完整个人数据流、CAS 真实身份接入保留；"
        "移动端键盘避让、320px 窄屏可用等细节经专项审计修复。",
    ]),
    ("三、技术架构", [
        ("h3", "3.1 总体架构：五层分工"),
        ("img", "1-architecture.png", "图 1  总体架构：五层结构与外部服务"),
        "五层清晰分工：前端（三工作区 SPA + 独立管理端）/ 应用层（FastAPI + SSE + ChatManager 运行治理："
        "队列、超时、取消、busy 限制）/ 智能体层（LangGraph 图 + 30+ 工具注册表）/ 服务层（CAS/jw/young/LLM 熔断）"
        "/ 数据层（三 SQLite 库 + ChromaDB/BM25 generation 版本化）。",
        "关键设计决策：同步 LangGraph 经有界线程池执行，不阻塞 FastAPI 事件循环；SSE 事件持久化 SQLite + "
        "进程内通知唤醒 + 1 秒轮询兜底（支持 Last-Event-ID 断线续传）；联网链路为博查搜索（国内直连）+ "
        "Crawl4AI 抓取 sidecar（egress/robots 健康检查前置），无自托管搜索引擎依赖；审核/发布 worker 与 Web 进程分离，"
        "发布约束（乐观锁、generation、不可变快照）保证知识库变更可追溯可回滚。",
        ("h3", "3.2 统一 QA 工作流：五角色协作"),
        ("img", "2-qa-workflow.png",
         "图 2  统一 QA 工作流：前置分流 → 确定性直连路由 / think 决策循环 → 并行工具 → 流式合成"),
        "五个智能体角色在一张 LangGraph 图内协作：意图分类器（本地 embedding）、决策器（think LLM，22 条规则约束）、"
        "工具执行器（act，单轮最多 3 个独立工具并行）、证据核验器（claim 级校验）、合成器（compose 流式）。",
        "编排方式是本项目的核心探索（对应设计取舍二）：规则与 LLM 自主决策的混合编排——高置信意图（成绩/课表/"
        "培养方案等）走确定性路由完全跳过决策 LLM（省 20–60 秒往返且行为可预期），长尾意图交给带规则约束的 LLM "
        "自主决策（≤4 轮循环，每轮工具表按意图裁剪）。",
    ]),
    ("四、功能模块", [
        ("h3", "4.1 数据源：全部为校内真实数据"),
        "综合教务 jw API（成绩/课表/考试）、青春科大 young 平台（第二课堂活动）、"
        "评课社区 icourse.club 公开数据（5667 个课程页 / 4.4 万条评论）、教务处官网与官方公众号"
        "（80 篇文档经采集审核入库）。覆盖教学（选课/成绩/课表）、生活（活动/日程/办事）、"
        "服务（官方入口/校园工具）三类场景。",
        "推荐语义经真实用户确认打磨：课程范围是硬条件，兴趣/工作量/教师/学期默认软排序，"
        "仅\"只要/必须\"升级为硬过滤——这来自多轮真实选课咨询的语义校准，而非想当然的过滤设计。",
        ("h3", "4.2 功能总表与模块清单"),
        ("table",
         ["场景", "用户问题示例", "小蜗的解决路径"],
         [
             ["选课决策", "\"推荐几门有意思的通识课\"", "培养方案定位 + 评课数据推荐 + 样本量披露 + 独立冲突检查"],
             ["成绩分析", "\"我的成绩怎么样\"", "直连成绩工具 → 数据卡 + 3 句要点（哪门课要小心）"],
             ["课表查询", "\"周四晚上有什么安排\"", "自然语言日期解析 → 日课表精确到节次时间"],
             ["办事流程", "\"学生证丢了怎么补办\"", "知识库命中官方文档 + 附官方 URL 可核验"],
             ["时效信息", "\"最新的选课通知\"", "联网检索（科大问题优先公众号）→ 证据门槛核验 → 附来源"],
             ["日程管理", "\"明天下午3点到4点加个提醒\"", "自然语言时间解析精确落点 + 冲突检测"],
             ["活动发现", "\"周末有什么活动可以报名\"", "青春科大实时数据 + 报名窗口/地点 + 每日个性化推荐"],
             ["强操作", "\"帮我退掉这门课\"", "如实说明无权代办 → 给出已核实官方入口 + 退课前压力评估"],
         ]),
        "模块清单：智能问答（RAG 混合检索）、联网证据通道、课业助手（成绩/GPA/周课表/考试/空教室）、"
        "今日卡弹窗、选课顾问（推荐/教师对比/冲突检查/退补选评估）、培养方案、日程与活动（自然语言时间/"
        "冲突检测/课表导入/四因子推荐）、官方入口与校园工具（申请—审核—上架）、生态工具（学生投稿）、"
        "管理后台（知识审核/工具治理/审计）。",
        "以上场景均已在公网环境实际运行验证（19 问实测 18 直接通过，答案正确），非纸面功能。",
        ("h3", "4.3 可扩展性与推广价值"),
        "生态工具协议：其他同学可以用\"Spec + 函数\"三步向小蜗投稿自制工具（eco: 前缀强制署名，"
        "坏 Spec 拒载不炸），智能体能力对学生社区开放扩展。",
        "校园服务与官方入口完全配置驱动（config/links.yaml），更换学校只需替换配置与知识库，"
        "架构（统一 QA 图/证据门槛/审核流水线）整体可迁移。",
        "工程实践可参考：225 项 Web 回归测试、260+ 工具断言、LLM 一致性金标评测的完整测试体系，"
        "对同类课程项目/比赛项目有直接复用价值。",
    ]),
    ("五、技术难度", [
        ("h3", "5.1 大模型应用深度（Prompt 工程 + RAG）"),
        "Prompt 工程不是零散技巧，而是成体系的约束设计。think 决策 Prompt 含 22 条规则约束，节选如下"
        "（源码 agents/qa/nodes.py，原文）：",
        ("pre",
         "8. 绝不编造数据：工具结果不足时继续 retrieve/call_tool/clarify，不要硬答\n"
         "9. 禁止重复调用：已有工具结果（status=done 的工具）不得再次调用同一工具，应转 compose 或 clarify\n"
         "13. 调用课程相关工具时，课程名关键词先解析为规范形式：补全常见简称（\"数分\"→\"数学分析\"…）\n"
         "14. 工具结果含 ambiguity=true 时：decision=clarify，引用候选反问用户选择哪个班型；禁止自行替用户做选择\n"
         "16. 个人数据工具调用时，args 必须携带 student_id（取自已提供的学生信息），不得省略\n"
         "18. 复合问题可并行调用多个相互独立的工具：单轮 tools 数组 1-3 个…工具间有先后依赖时必须分开轮次\n"
         "20. 生态工具结果转述时必须标注提供者署名与\"仅供参考\"，不得与官方数据混写\n"
         "21. 强操作无权代办→render_link(scene) 给官方入口…URL 只能来自 render_link 或知识库，禁止拼造"),
        "compose 合成 Prompt 同样以反幻觉约束为核心，节选：",
        ("pre",
         "- 课程学分、均分、样本量、学期等数值必须取自工具返回结果，不得猜测、修改或补充\n"
         "- 当官方信息含\"因专业/因年份/因人群而有差异的多个数值\"时，必须区分适用对象作答，\n"
         "  不得用其中某一个特例数值代表整体\n"
         "- 不得提及未通过工具实际查询到的数据…工具未查过的一律不主动提及\n"
         "- 数据不足时如实说明并引导用户补充信息，不编造\n"
         "- source=generic 时必须醒目说明\"专业通用参考，不是个人培养方案\""),
        "这些规则全部来自真实测试中发现的问题（开发日志可查）：规则 8/9 源于工具重复调用 4 次的实测 bug，"
        "规则 14 源于多班型课程答错班的实测案例，\"特例不得代表整体\"源于学费问题把传播学学费答成全校学费的"
        "实测错误——每条规则都对应一个被修复的真实缺陷，这是 Prompt 工程深度的最好证明。",
        "RAG 超出基础问答的完整工程闭环：",
        ("img", "3-evidence.png", "图 3  联网证据流水线：隐私清洗 → 多通道检索 → 证据门槛判定 → 知识治理旁路"),
        "混合检索（ChromaDB 向量 + BM25，维度不匹配自动降级）；查询改写与 1–3 个并列子查询"
        "（\"退学和休学分别怎么办\"拆分检索）；语义缓存（命中复用 + 防陈旧护栏）；"
        "claim 级证据校验（引用不通过则该条降级 insufficient 并改写，不得输出原事实正文）；"
        "证据门槛（一手权威来源或双独立一致来源才输出确定结论）；"
        "知识治理流水线（采集→LLM 清洗→逐块审核→generation 发布→双索引重建，不可变快照可回滚）。",
        ("h3", "5.2 调试与优化能力（输出可靠性与响应效率）"),
        ("img", "4-streaming.png", "图 4  流式响应时序（真实 SSE 事件时间线，问题\"我的成绩怎么样\"）"),
        "性能优化是一整条链路，而非单点：意图裁剪工具表 → think 提示词精简 → 单轮并行工具调用 → "
        "本地+联网双通道并行 → compose 增量流式 → 超时部分收尾。实测数据：",
        ("table",
         ["指标", "优化前（实测）", "优化后（实测）"],
         [
             ["个人数据问答首字延迟", "先经 think LLM 决策往返 20–60 秒，再整篇等待生成", "确定性直连路由，首字 3.1 秒"],
             ["正文呈现方式", "整篇生成完毕一次性返回", "增量流式推送（示例问题 21 段，~16 字/批）"],
             ["生成超时行为", "整体失败（UPSTREAM_TIMEOUT）", "保留已流出正文部分收尾（truncated 标注）"],
             ["LLM 平台故障", "问答全不可用", "三层降级：意图本地分类 / 工具摘要直出 / 熔断不传染"],
             ["回归测试", "—", "pytest 225 项 / 前端 22 项 / verify 系列 260+ 断言 / LLM 金标 12/12"],
         ]),
        "可靠性同样体系化：LLM 熔断窗（仅连接级失败开窗，读超时不动窗）；证据抽取与就绪探针多级重试；数据降级带来源标注；"
        "并发有界队列与 503 RUN_BUSY 背压；prompt 注入快照测试（qa_consistency 断言体系）。"
        "安全工程贯穿全链路：SSRF 防护（内网/回环/云元数据/重定向）、身份边界（§2.2）、demo/生产数据命名空间隔离、"
        "审计不可变、密钥经 .env 注入不入 git。",
    ]),
    ("六、运行环境与演示", [
        "生产部署：Linux 服务器（Ubuntu 24.04）完整运行——Web:8000（公网 8850 转发）/ 审核发布 worker 常驻 / "
        "Streamlit 回退入口；公网地址 http://114.214.241.119:8850 当前可访问，readiness 健康端点四项哨兵"
        "（数据库/审核库/批准索引/检索质量）全绿。",
        "测试基线：pytest tests/web 225 项 / 前端 vitest 22 项 / verify 系列 260+ 断言 / LLM 一致性金标 12/12。",
        "文档体系：仓库维护 README（快速上手）、PRODUCT/DESIGN（产品与设计系统事实）、tool-specs（30+ 工具详细规格）、"
        "Web 技术规格（Q1–Q80 已确认契约）、部署规格与运维记录、开发日志；模型调用说明见作品简介第五节"
        "（deepseek-v4-flash/qwen3-embedding/unlimited-ocr/博查搜索的用途与调用方式）。",
        "演示：≤5 分钟视频按\"工作流总览 → 问答与来源核验 → 个人数据直连 → 选课顾问 → 学业工作区 → "
        "联网证据门槛 → 校园服务与知识治理\"编排；公网演示环境可供评委实际操作验证（合成演示身份，"
        "个人数据明确标识为演示数据）。",
    ]),
]


# ============================ 渲染 ============================

def _set_font(run, size: float, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _label_split(text: str) -> tuple[str, str] | None:
    """「标签：内容」格式的段落，标签加粗（无标签返回 None）。"""
    for sep in ("：", ":"):
        idx = text.find(sep)
        if 0 < idx <= 16 and not text.startswith(("（", "（")):
            return text[: idx + 1], text[idx + 1:]
    return None


def build_docx(title: str, subtitle: str, sections: list, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(title), 18, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(subtitle), 11)

    for heading, items in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        _set_font(p.add_run(heading), 14, bold=True, color=PRIMARY)
        for item in items:
            if isinstance(item, str):
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.line_spacing = 1.3
                label = _label_split(item)
                if label:
                    _set_font(para.add_run(label[0]), 10.5, bold=True)
                    _set_font(para.add_run(label[1]), 10.5)
                else:
                    _set_font(para.add_run(item), 10.5)
            elif item[0] == "h3":
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(8)
                _set_font(para.add_run(item[1]), 12, bold=True)
            elif item[0] == "img":
                _, filename, caption = item
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(str(DIAG_DIR / filename), width=Inches(6.1))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_font(cap.add_run(caption), 9, color=MUTED)
            elif item[0] == "table":
                _, headers, rows = item
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                for j, htext in enumerate(headers):
                    cell_p = table.rows[0].cells[j].paragraphs[0]
                    _set_font(cell_p.add_run(htext), 9.5, bold=True)
                for i, row in enumerate(rows):
                    for j, ctext in enumerate(row):
                        cell_p = table.rows[i + 1].cells[j].paragraphs[0]
                        _set_font(cell_p.add_run(ctext), 9)
                doc.add_paragraph()
            elif item[0] == "pre":
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(item[1])
                run.font.size = Pt(9)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.color.rgb = RGBColor(0x17, 0x24, 0x2D)

    doc.save(out_path)


HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8">
<title>{title}</title>
<style>
  @page {{ size: A4; margin: 18mm 16mm; }}
  body {{ font-family: "Microsoft YaHei", "PingFang SC", sans-serif; color: #1c2430;
         font-size: 11pt; line-height: 1.65; max-width: 210mm; margin: 0 auto; }}
  h1 {{ color: #034ea1; font-size: 20pt; text-align: center; margin: 0 0 4pt; }}
  .subtitle {{ text-align: center; color: #566672; font-size: 11pt; margin-bottom: 18pt; }}
  h2 {{ color: #034ea1; font-size: 13.5pt; margin: 16pt 0 6pt;
       border-left: 4pt solid #034ea1; padding-left: 8pt; }}
  h3 {{ color: #1c2430; font-size: 12pt; margin: 12pt 0 4pt; }}
  p {{ margin: 0 0 7pt; text-align: justify; }}
  p.lbl {{ margin: 8pt 0 2pt; }}
  p.lbl b {{ color: #034ea1; }}
  img {{ max-width: 178mm; display: block; margin: 8pt auto 2pt; }}
  .caption {{ text-align: center; color: #566672; font-size: 9pt; margin: 0 0 10pt; }}
  table {{ border-collapse: collapse; width: 100%; margin: 6pt 0 10pt; font-size: 9.5pt; }}
  th {{ background: #e2ecf7; color: #1c2430; border: 0.5pt solid #aebbc4; padding: 4pt 6pt; text-align: left; }}
  td {{ border: 0.5pt solid #aebbc4; padding: 4pt 6pt; }}
  pre {{ background: #f3f6fb; border-left: 3pt solid #034ea1; padding: 7pt 10pt;
        font-family: Consolas, monospace; font-size: 9pt; line-height: 1.5;
        white-space: pre-wrap; margin: 4pt 0 10pt; }}
  .footer {{ margin-top: 24pt; color: #566672; font-size: 9pt; border-top: 0.5pt solid #d5dde6; padding-top: 6pt; }}
</style></head>
<body>
<h1>{title}</h1>
<div class="subtitle">{subtitle}</div>
{body}
<div class="footer">小蜗（Xiaowo）· 一〇七杯算力与智能体开发大赛 · 智能体赛道本科生队伍 · 2026-09</div>
</body></html>
"""


def _esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _para_html(text: str) -> str:
    label = _label_split(text)
    if label:
        return f'<p class="lbl"><b>{_esc(label[0])}</b>{_esc(label[1])}</p>'
    return f"<p>{_esc(text)}</p>"


def build_html(title: str, subtitle: str, sections: list, out_path: Path) -> None:
    parts: list[str] = []
    for heading, items in sections:
        parts.append(f"<h2>{_esc(heading)}</h2>")
        for item in items:
            if isinstance(item, str):
                parts.append(_para_html(item))
            elif item[0] == "h3":
                parts.append(f"<h3>{_esc(item[1])}</h3>")
            elif item[0] == "img":
                _, filename, caption = item
                src = (DIAG_DIR / filename).as_uri()
                parts.append(f'<img src="{src}" alt="{_esc(caption)}">')
                parts.append(f'<div class="caption">{_esc(caption)}</div>')
            elif item[0] == "table":
                _, headers, rows = item
                cells = "".join(f"<th>{_esc(h)}</th>" for h in headers)
                body = "".join(
                    "<tr>" + "".join(f"<td>{_esc(c)}</td>" for c in row) + "</tr>" for row in rows
                )
                parts.append(f"<table><tr>{cells}</tr>{body}</table>")
            elif item[0] == "pre":
                parts.append(f"<pre>{_esc(item[1])}</pre>")
    html = HTML_TEMPLATE.format(
        title=_esc(title), subtitle=_esc(subtitle), body="\n".join(parts)
    )
    out_path.write_text(html, encoding="utf-8")


def main() -> None:
    for name, title, subtitle, sections in (
        ("作品简介_小蜗", BRIEF_TITLE, BRIEF_SUBTITLE, BRIEF_SECTIONS),
        ("设计文档_小蜗", DESIGN_TITLE, DESIGN_SUBTITLE, DESIGN_SECTIONS),
    ):
        build_docx(title, subtitle, sections, OUT_DIR / f"{name}.docx")
        build_html(title, subtitle, sections, OUT_DIR / f"{name}.html")
        print(f"生成 {name}.docx / .html")


if __name__ == "__main__":
    main()
